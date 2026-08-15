"""Focused regression tests for deterministic person.id derivation.

person.id must be derived by CareerOS from the extracted person identity
rather than trusted from the LLM, so a low-quality extraction (e.g. an id
copied from the few-shot prompt example, like ``person-smith``) can never
override the deterministic id.
"""

import json
from pathlib import Path
from typing import Any

import pytest
import yaml

from careeros.acquisition.llm_extractor import LLMExtractor
from careeros.acquisition.person_data import PersonData
from careeros.acquisition.pipeline import AcquisitionPipeline
from careeros.acquisition.profile_builder import CanonicalProfileBuilder
from careeros.acquisition.utils import person_id_from_name
from careeros.ai.mock_provider import MockAIProvider


def test_person_id_from_name_normal() -> None:
    assert person_id_from_name("Raul Gongora Betancourt") == "person-raul-gongora-betancourt"


def test_person_id_from_name_unicode_folds_to_ascii() -> None:
    assert person_id_from_name("Raúl González Betancourt") == "person-raul-gonzalez-betancourt"
    assert person_id_from_name("Åsa Lindqvist") == "person-asa-lindqvist"
    assert person_id_from_name("José María") == "person-jose-maria"


def test_person_id_from_name_empty_and_garbage() -> None:
    assert person_id_from_name("") == ""
    assert person_id_from_name("   ") == ""
    assert person_id_from_name("!!!") == ""


def test_person_id_from_name_is_deterministic() -> None:
    name = "Raul Gongora Betancourt"
    assert person_id_from_name(name) == person_id_from_name(name)


def _build_person_id(person: PersonData) -> str:
    profile = CanonicalProfileBuilder().build(person)
    return profile["person"]["id"]


def test_build_derives_id_from_name_ignoring_llm_id() -> None:
    person = PersonData(
        id="person-smith",
        first_name="Raul",
        last_name="Gongora",
        full_name="Raul Gongora Betancourt",
    )
    assert _build_person_id(person) == "person-raul-gongora-betancourt"


def test_build_uses_first_last_when_full_name_missing() -> None:
    person = PersonData(
        id="person-smith",
        first_name="Anna",
        last_name="Lindqvist",
        full_name="",
    )
    assert _build_person_id(person) == "person-anna-lindqvist"


def test_build_preserves_fallback_when_no_usable_name() -> None:
    person = PersonData(id="person-llm", first_name="", last_name="", full_name="")
    assert _build_person_id(person) == "person-llm"

    anonymous = PersonData(id="", first_name="", last_name="", full_name="")
    assert _build_person_id(anonymous) == "person-unknown"


def _create_test_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Raul Gongora Betancourt", level=1)
    doc.add_paragraph("Senior IT DevSecOps Specialist")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Qred Bank - Senior DevSecOps Engineer (2022-2025)")
    doc.save(str(path))


def _run_pipeline_with(docx_path: Path, extractor: LLMExtractor, out: Path) -> dict:
    pipeline = AcquisitionPipeline(llm_extractor=extractor)
    result = pipeline.run(docx_path, out)
    assert result == out.resolve()
    return yaml.safe_load(out.read_text(encoding="utf-8"))


def _extractor_for(payload: dict[str, Any]) -> LLMExtractor:
    def generator(prompt: str, temperature: float, timeout: float) -> str:
        return json.dumps(payload, ensure_ascii=False)

    return LLMExtractor(provider=MockAIProvider(generator=generator))


def test_pipeline_llm_person_smith_cannot_override_raul(tmp_path: Path) -> None:
    """Regression: LLM copies person-smith from the prompt; deterministic id must win."""
    docx = tmp_path / "raul-gongora-cv.docx"
    _create_test_docx(docx)
    payload = {
        "person": {
            "id": "person-smith",
            "firstName": "Raul",
            "lastName": "Gongora",
            "fullName": "Raul Gongora Betancourt",
        },
        "experiences": [
            {"id": "exp-qred", "organization": "Qred Bank", "title": "Senior DevSecOps Engineer"}
        ],
        "skills": [],
        "education": [],
    }
    loaded = _run_pipeline_with(docx, _extractor_for(payload), tmp_path / "out.yaml")

    assert loaded["person"]["id"] == "person-raul-gongora-betancourt"
    assert loaded["person"]["names"][0]["value"] == "Raul Gongora Betancourt"


def test_pipeline_preserves_fallback_when_no_usable_name(tmp_path: Path) -> None:
    """No usable name: the existing person-unknown fallback is preserved."""
    docx = tmp_path / "cv.docx"
    _create_test_docx(docx)
    payload = {
        "person": {"id": None, "firstName": "", "lastName": "", "fullName": ""},
        "experiences": [],
        "skills": [{"name": "Python"}],
        "education": [],
    }
    loaded = _run_pipeline_with(docx, _extractor_for(payload), tmp_path / "out.yaml")

    assert loaded["person"]["id"] == "person-unknown"
