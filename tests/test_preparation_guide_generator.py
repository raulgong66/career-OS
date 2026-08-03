"""M1.16 — Interview Preparation Guide generation tests.

Covers Markdown/DOCX generation from InterviewPlan through the generic
generator registry pipeline, section ordering, question grouping, suggested
answers, deterministic output, empty plan, edge cases, generator registration,
artifact lifecycle compatibility, and inference.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
import yaml
from docx import Document

from careeros import (
    InterviewEngine,
    InterviewPlan,
)
from careeros.exceptions import ValidationError
from careeros.export_contract import ExportContract, ExportContractBuilder
from careeros.generators import (
    DocxPreparationGuideGenerator,
    GeneratorRegistry,
    MarkdownPreparationGuideGenerator,
    default_generator_registry,
)
from careeros.pipelines import _infer_artifact_type, generate_artifact
from careeros.resolution import apply_resolution
from careeros.schema_loader import SchemaLoader


# --------------------------------------------------------------------------
# Fixtures
# --------------------------------------------------------------------------


def make_profile() -> dict:
    """Profile with rich cross-linked evidence matching interview engine tests."""
    return {
        "profileVersion": "1.0.0",
        "person": {
            "id": "person-tester",
            "names": [{"value": "Test Person", "usage": "professional"}],
            "positioning": {"headline": "Senior Platform Engineer"},
        },
        "organizations": [
            {"id": "org-acme", "name": "ACME Corp"},
            {"id": "org-beta", "name": "Beta Labs"},
        ],
        "experiences": [
            {
                "id": "exp-acme",
                "title": "Platform Engineer",
                "organizationRefs": [{"id": "org-acme", "type": "organization"}],
                "scope": (
                    "Led migration of 200 services to Kubernetes and AWS,"
                    " reducing infrastructure costs by 40%."
                ),
            },
            {
                "id": "exp-beta",
                "title": "Founder",
                "organizationRefs": [{"id": "org-beta", "type": "organization"}],
                "scope": (
                    "Founded a startup and built the product end to end"
                    " with Python and Django."
                ),
            },
        ],
        "projects": [
            {
                "id": "project-platform",
                "name": "Platform Reboot",
                "description": (
                    "Rebuilt the deployment platform on Kubernetes with GitOps."
                ),
            }
        ],
        "skills": [
            {
                "id": "skill-kubernetes",
                "name": "Kubernetes",
                "description": "AWS, GCP, container orchestration, helm",
                "extensions": {"experienceEvidence": [{"experienceId": "exp-acme"}]},
            },
            {
                "id": "skill-python",
                "name": "Python",
                "description": "Django, REST APIs, automation",
                "extensions": {"experienceEvidence": [{"experienceId": "exp-beta"}]},
            },
        ],
        "achievements": [
            {
                "id": "achievement-cost-reduction",
                "statement": (
                    "Reduced infrastructure costs by 40% through the"
                    " Kubernetes migration."
                ),
                "contextRefs": [{"id": "exp-acme", "type": "experience"}],
            }
        ],
        "professionalSummaries": [
            {
                "id": "summary-1",
                "text": "Senior platform engineer focused on Kubernetes and automation.",
            }
        ],
        "targetContexts": [
            {
                "id": "ctx-devsecops",
                "label": "DevSecOps target",
                "role": "Senior DevSecOps Specialist",
            }
        ],
        "artifacts": [
            {
                "id": "artifact-interview-guide",
                "title": "Interview Preparation Guide",
                "artifactType": "INTERVIEW_PREPARATION_GUIDE",
                "targetContextRefs": [
                    {"id": "ctx-devsecops", "type": "targetContext"}
                ],
                "sourceRefs": [
                    {"id": "summary-1", "type": "professional_summary"},
                    {"id": "exp-acme", "type": "experience"},
                    {"id": "exp-beta", "type": "experience"},
                    {"id": "project-platform", "type": "project"},
                    {"id": "skill-kubernetes", "type": "skill"},
                    {"id": "skill-python", "type": "skill"},
                    {"id": "achievement-cost-reduction", "type": "achievement"},
                ],
            }
        ],
        "education": [],
        "certifications": [],
    }


@pytest.fixture
def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


@pytest.fixture
def profile() -> dict:
    return make_profile()


@pytest.fixture
def schema_loader(repo_root: Path) -> SchemaLoader:
    return SchemaLoader(repo_root / "schemas")


def _generate_guide(
    generator, profile: dict, repo_root: Path, schema_loader: SchemaLoader
) -> str:
    """Build an export contract with an attached interview plan and generate."""
    from careeros.interview.engine import build_preparation_plan

    contract = ExportContractBuilder(schema_loader).build(
        profile, "artifact-interview-guide", validate=False
    )
    plan = build_preparation_plan(profile, target_contexts=contract.target_contexts)
    contract.interview_plan = plan
    return generator.generate(contract)


# --------------------------------------------------------------------------
# Markdown generation
# --------------------------------------------------------------------------


def test_markdown_renders_expected_sections(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "# Interview Preparation Guide" in output
    assert "**Target Role:** Senior DevSecOps Specialist" in output
    assert "## Candidate Summary" in output
    assert "## Technical Questions" in output
    assert "## Behavioral Questions" in output
    assert "## Leadership Questions" in output
    assert "## Project Deep Dive" in output
    assert "## Problem Solving" in output
    assert "## Career Motivation" in output
    assert "## Suggested STAR Outlines" in output
    assert "## Preparation Checklist" in output
    assert "## Evidence Notes" in output


def test_section_ordering(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    sections = [line for line in output.splitlines() if line.startswith("## ")]
    expected_order = [
        "Candidate Summary",
        "Technical Questions",
        "Behavioral Questions",
        "Leadership Questions",
        "Project Deep Dive",
        "Problem Solving",
        "Career Motivation",
        "Suggested STAR Outlines",
        "Preparation Checklist",
        "Evidence Notes",
    ]
    found = [s.replace("## ", "") for s in sections]
    assert found == expected_order


def test_question_grouping_by_category(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    sections = output.split("## ")
    technical = next(s for s in sections if s.startswith("Technical Questions\n"))
    assert "Kubernetes" in technical
    assert "Python" in technical
    leadership = next(s for s in sections if s.startswith("Leadership Questions\n"))
    assert "Founder" in leadership


def test_suggested_answer_renders_star_components(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "- Situation:" in output
    assert "- Action:" in output
    assert "- Result:" in output
    assert "- Evidence:" in output


def test_omits_empty_answer_components(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    for line in output.splitlines():
        if line.startswith("- Situation:"):
            assert line.strip() != "- Situation:"
        if line.startswith("- Task:"):
            assert line.strip() != "- Task:"


def test_deterministic_output(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    gen = MarkdownPreparationGuideGenerator()
    first = _generate_guide(gen, profile, repo_root, schema_loader)
    second = _generate_guide(gen, profile, repo_root, schema_loader)
    assert first == second


def test_empty_plan_renders_minimal_document() -> None:
    gen = MarkdownPreparationGuideGenerator()
    empty_profile = {
        "profileVersion": "1.0.0",
        "person": {"id": "person-empty"},
        "skills": [],
        "experiences": [],
        "projects": [],
        "achievements": [],
    }
    contract = ExportContract(
        profile_version="1.0.0",
        artifact_id="guide-1",
        artifact_type="INTERVIEW_PREPARATION_GUIDE",
        person=empty_profile["person"],
        artifact={"id": "guide-1", "title": "Guide"},
    )
    contract.interview_plan = InterviewEngine().generate_plan(empty_profile)
    output = gen.generate(contract)
    assert output.strip()
    assert "## Technical Questions" not in output
    assert "## Evidence Notes" not in output


def test_no_duplicated_questions(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    q_bullets = [line for line in output.splitlines() if line.startswith("• ")]
    assert len(q_bullets) == len(set(q_bullets))


def test_no_internal_ids_exposed(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    for phrase in ("q-technical-", "exp-acme", "skill-kubernetes", "evidenceRefs",
                   "contextRefs", "person-tester"):
        assert phrase not in output, f"Leaked internal id: {phrase}"


def test_target_role_rendered_when_set(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "**Target Role:** Senior DevSecOps Specialist" in output


def test_target_role_omitted_when_none(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    profile["targetContexts"] = []
    profile["artifacts"][0]["targetContextRefs"] = []
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "**Target Role:**" not in output


def test_star_outlines_link_questions(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "- Technical 1:" in output
    assert "- Behavioral 1:" in output or "- Behavioral 2:" in output


def test_evidence_notes_omit_internal_ids(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    notes_start = output.index("## Evidence Notes")
    notes_text = output[notes_start:]
    for id_pat in ("exp-acme", "skill-kubernetes", "q-", "person-tester"):
        assert id_pat not in notes_text, f"Evidence notes leaked: {id_pat}"


def test_evidence_notes_grouped_by_element_type(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "- Skills:" in output
    assert "- Experiences:" in output


def test_rejects_non_interview_contract() -> None:
    gen = MarkdownPreparationGuideGenerator()
    plan = InterviewEngine().generate_plan(make_profile())
    contract = ExportContract(
        profile_version="1.0.0",
        artifact_id="cv-1",
        artifact_type="CV",
        person={},
        artifact={"id": "cv-1", "title": "CV"},
        interview_plan=plan,
    )
    with pytest.raises(ValidationError, match="Unsupported artifact type"):
        gen.generate(contract)


def test_requires_interview_plan(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    gen = MarkdownPreparationGuideGenerator()
    contract = ExportContractBuilder(schema_loader).build(
        profile, "artifact-interview-guide", validate=False
    )
    with pytest.raises(ValidationError, match="requires an interview_plan"):
        gen.generate(contract)


def test_preparation_checklist_present(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "## Preparation Checklist" in output
    assert "STAR format" in output
    assert "structured answers" in output


def test_candidate_summary_from_profile(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        MarkdownPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert "Senior platform engineer" in output


# --------------------------------------------------------------------------
# DOCX generation
# --------------------------------------------------------------------------


def test_docx_generation_returns_bytes(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        DocxPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    assert isinstance(output, bytes)
    assert output.startswith(b"PK")


def test_docx_renders_expected_content(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        DocxPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    document = Document(BytesIO(output))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Interview Preparation Guide" in text
    assert "Senior DevSecOps Specialist" in text
    assert "Candidate Summary" in text
    assert "Technical Questions" in text
    assert "Preparation Checklist" in text
    assert "Evidence Notes" in text


def test_docx_no_internal_ids_in_text(
    repo_root: Path, profile: dict, schema_loader: SchemaLoader
) -> None:
    output = _generate_guide(
        DocxPreparationGuideGenerator(), profile, repo_root, schema_loader
    )
    document = Document(BytesIO(output))
    text = "\n".join(p.text for p in document.paragraphs)
    for id_val in ("exp-acme", "skill-kubernetes", "evidenceRefs"):
        assert id_val not in text, f"DOCX leaked: {id_val}"


def test_docx_rejects_non_interview_contract() -> None:
    gen = DocxPreparationGuideGenerator()
    plan = InterviewEngine().generate_plan(make_profile())
    contract = ExportContract(
        profile_version="1.0.0",
        artifact_id="cv-1",
        artifact_type="CV",
        person={},
        artifact={"id": "cv-1", "title": "CV"},
        interview_plan=plan,
    )
    with pytest.raises(ValidationError, match="Unsupported artifact type"):
        gen.generate(contract)


# --------------------------------------------------------------------------
# Registry
# --------------------------------------------------------------------------


def test_default_registry_registers_markdown_preparation_guide_generator() -> None:
    gen = default_generator_registry().resolve("INTERVIEW_PREPARATION_GUIDE", "markdown")
    assert isinstance(gen, MarkdownPreparationGuideGenerator)


def test_default_registry_registers_docx_preparation_guide_generator() -> None:
    gen = default_generator_registry().resolve("INTERVIEW_PREPARATION_GUIDE", "docx")
    assert isinstance(gen, DocxPreparationGuideGenerator)


# --------------------------------------------------------------------------
# Generation pipeline (full integration)
# --------------------------------------------------------------------------


class CapturingGenerator:
    def __init__(self) -> None:
        self.captured_contract: ExportContract | None = None

    def generate(self, contract: ExportContract) -> str:
        self.captured_contract = contract
        return "captured"


def test_pipeline_attaches_interview_plan_to_contract(
    repo_root: Path, tmp_path: Path
) -> None:
    profile = make_profile()
    profile_file = tmp_path / "profile.yaml"
    profile_file.write_text(yaml.safe_dump(profile), encoding="utf-8")

    registry = GeneratorRegistry()
    capturer = CapturingGenerator()
    registry.register("INTERVIEW_PREPARATION_GUIDE", "text", capturer)

    generate_artifact(
        profile_file,
        "artifact-interview-guide",
        "text",
        SchemaLoader(repo_root / "schemas"),
        registry,
    )
    assert capturer.captured_contract is not None
    plan = capturer.captured_contract.interview_plan
    assert plan is not None
    assert isinstance(plan, InterviewPlan)
    assert plan.target_role == "Senior DevSecOps Specialist"
    assert plan.question_count > 0


def test_pipeline_generates_markdown_guide_via_registry(
    repo_root: Path, tmp_path: Path
) -> None:
    profile = make_profile()
    profile_file = tmp_path / "profile.yaml"
    profile_file.write_text(yaml.safe_dump(profile), encoding="utf-8")

    output = generate_artifact(
        profile_file,
        "artifact-interview-guide",
        "markdown",
        SchemaLoader(repo_root / "schemas"),
    )
    assert "# Interview Preparation Guide" in output
    assert "## Technical Questions" in output


def test_pipeline_generates_docx_guide_via_registry(
    repo_root: Path, tmp_path: Path
) -> None:
    profile = make_profile()
    profile_file = tmp_path / "profile.yaml"
    profile_file.write_text(yaml.safe_dump(profile), encoding="utf-8")

    output = generate_artifact(
        profile_file,
        "artifact-interview-guide",
        "docx",
        SchemaLoader(repo_root / "schemas"),
    )
    assert isinstance(output, bytes)
    assert output.startswith(b"PK")
    document = Document(BytesIO(output))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Interview Preparation Guide" in text


# --------------------------------------------------------------------------
# Artifact lifecycle (resolution → staleness)
# --------------------------------------------------------------------------


def test_interview_guide_artifact_in_export_contract(
    repo_root: Path, schema_loader: SchemaLoader
) -> None:
    profile = make_profile()
    contract = ExportContractBuilder(schema_loader).build(
        profile, "artifact-interview-guide", validate=False
    )
    assert contract.artifact_type == "INTERVIEW_PREPARATION_GUIDE"
    assert contract.artifact_id == "artifact-interview-guide"


def test_interview_guide_artifact_has_source_refs(
    repo_root: Path, schema_loader: SchemaLoader
) -> None:
    profile = make_profile()
    contract = ExportContractBuilder(schema_loader).build(
        profile, "artifact-interview-guide", validate=False
    )
    assert len(contract.sources) > 0
    source_types = {s.type for s in contract.sources}
    assert "experience" in source_types
    assert "skill" in source_types


def test_resolution_marks_interview_guide_stale() -> None:
    import copy
    profile = copy.deepcopy(make_profile())
    apply_resolution(
        profile,
        triggered_rule="ExperienceNoTechnologiesRule",
        element_id="exp-acme",
        technologies=["Docker", "K8s"],
    )
    guide = next(
        a for a in profile["artifacts"] if a["id"] == "artifact-interview-guide"
    )
    assert guide.get("status") == "stale"


# --------------------------------------------------------------------------
# Inference helpers
# --------------------------------------------------------------------------


def test_infer_artifact_type_from_id_contains_interview() -> None:
    assert _infer_artifact_type({"id": "guide-interview", "title": "G"}) == "INTERVIEW_PREPARATION_GUIDE"


def test_infer_artifact_type_from_title_contains_interview() -> None:
    assert _infer_artifact_type({"id": "g", "title": "Interview Prep Guide"}) == "INTERVIEW_PREPARATION_GUIDE"


def test_infer_artifact_type_from_explicit_type() -> None:
    assert _infer_artifact_type(
        {"id": "a", "title": "G", "artifactType": "INTERVIEW_PREPARATION_GUIDE"}
    ) == "INTERVIEW_PREPARATION_GUIDE"


# --------------------------------------------------------------------------
# build_preparation_plan public API
# --------------------------------------------------------------------------


def test_build_preparation_plan_exported_from_facade() -> None:
    from careeros import build_preparation_plan

    assert build_preparation_plan is not None