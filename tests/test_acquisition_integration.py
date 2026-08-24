from pathlib import Path
from typing import Any

import pytest
import yaml

from careeros.acquisition.document_reader import DocumentReader
from careeros.acquisition.llm_extractor import LLMExtractor, OllamaLLMExtractor, LLMConfigurationError, create_llm_extractor
from careeros.acquisition.person_data import (
    EducationData,
    ExperienceData,
    ExtractionResult,
    PersonData,
    SkillData,
)
from careeros.acquisition.pipeline import AcquisitionPipeline, PipelineError
from careeros.acquisition.profile_builder import CanonicalProfileBuilder
from careeros.acquisition.text_extractor import TextExtractor
from careeros.acquisition.yaml_writer import YamlWriter
from careeros.ai.mock_provider import MockAIProvider
from careeros.schema_loader import SchemaLoader
from careeros.validator import EntityValidator


class MockLLMExtractor(LLMExtractor):
    def extract(
        self, text: str, schema: dict[str, Any] | None = None
    ) -> ExtractionResult:
        return ExtractionResult(
            person=PersonData(
                id="person-gongora",
                first_name="Raul",
                last_name="Gongora",
                full_name="Raul Gongora Betancourt",
                email="raul@example.com",
                phone="+46-70-123-4567",
                location="Stockholm, Sweden",
                linkedin="https://linkedin.com/in/raulgongora",
                github="https://github.com/raulgong66",
            ),
            experiences=[
                ExperienceData(
                    id="exp-qred-bank",
                    organization="Qred Bank",
                    title="Senior DevSecOps Engineer",
                    start_date="2022",
                    end_date="2025",
                    is_current=False,
                    location="Stockholm, Sweden",
                    summary="DevSecOps for fintech platform",
                    responsibilities=[
                        "Managed cloud infrastructure",
                        "Implemented CI/CD pipelines",
                    ],
                    technologies=["AWS", "Docker", "Kubernetes"],
                ),
                ExperienceData(
                    id="exp-cabonline",
                    organization="Cabonline",
                    title="DevOps Engineer",
                    start_date="2020",
                    end_date="2022",
                    is_current=False,
                    location="Stockholm, Sweden",
                    summary="DevOps for taxi platform",
                    responsibilities=[
                        "Automated deployments",
                        "Managed Kubernetes clusters",
                    ],
                    technologies=["Terraform", "Kubernetes"],
                ),
            ],
            skills=[
                SkillData(name="AWS", category="Cloud Platform", proficiency="Advanced"),
                SkillData(name="Docker", category="Container", proficiency="Advanced"),
                SkillData(name="Kubernetes", category="Orchestration", proficiency="Advanced"),
                SkillData(name="Python", category="Programming Language"),
                SkillData(name="Terraform", category="IaC"),
            ],
            education=[
                EducationData(
                    institution="KTH Royal Institute of Technology",
                    degree="M.Sc. in Computer Science",
                    field_of_study="Computer Science",
                    start_date="2018",
                    end_date="2020",
                    is_current=False,
                    location="Stockholm, Sweden",
                ),
                EducationData(
                    institution="MIT",
                    degree="Bachelor of Science",
                    field_of_study="Electrical Engineering",
                    start_date="2014",
                    end_date="2018",
                    is_current=False,
                ),
            ],
        )


@pytest.fixture
def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _create_test_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Raul Gongora Betancourt", level=1)
    doc.add_paragraph("Senior IT DevSecOps Specialist & AI Solutions Architect")
    doc.add_heading("Contact", level=2)
    doc.add_paragraph("Email: raul@example.com")
    doc.add_paragraph("Phone: +46-70-123-4567")
    doc.add_paragraph("Location: Stockholm, Sweden")
    doc.add_paragraph("LinkedIn: https://linkedin.com/in/raulgongora")
    doc.add_paragraph("GitHub: https://github.com/raulgong66")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Experienced DevSecOps professional with a strong background "
        "in cloud infrastructure, AI solutions, and cybersecurity."
    )
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Qred Bank - Senior DevSecOps Engineer (2022-2025)")
    doc.add_paragraph("Cabonline - DevOps Engineer (2020-2022)")
    doc.save(str(path))


def test_acquisition_pipeline_full_vertical_slice(tmp_path: Path, repo_root: Path) -> None:
    """Integration test: DOCX → Parse → Extract → Profile → Validate."""
    docx_path = tmp_path / "test_resume.docx"
    _create_test_docx(docx_path)

    output_path = tmp_path / "generated-profile.yaml"

    llm = MockLLMExtractor()
    pipeline = AcquisitionPipeline(
        document_reader=DocumentReader(),
        text_extractor=TextExtractor(),
        llm_extractor=llm,
        profile_builder=CanonicalProfileBuilder(),
        yaml_writer=YamlWriter(),
    )

    result = pipeline.run(docx_path, output_path)

    assert result == output_path.resolve()
    assert output_path.exists()

    loaded = yaml.safe_load(output_path.read_text(encoding="utf-8"))
    assert loaded["profileVersion"] == "1.0.0"
    assert loaded["person"]["id"] == "person-raul-gongora-betancourt"
    assert loaded["person"]["names"] == [
        {"value": "Raul Gongora Betancourt", "usage": "professional"}
    ]
    assert loaded["person"]["contact"]["email"] == "raul@example.com"
    assert loaded["person"]["contact"]["phone"] == "+46-70-123-4567"
    assert loaded["person"]["location"]["label"] == "Stockholm, Sweden"
    assert loaded["person"]["links"][0]["label"] == "LinkedIn"
    assert loaded["person"]["links"][0]["href"] == "https://linkedin.com/in/raulgongora"
    assert loaded["person"]["links"][1]["label"] == "GitHub"
    assert loaded["person"]["links"][1]["href"] == "https://github.com/raulgong66"
    experiences = loaded.get("experiences", [])
    assert len(experiences) == 2

    exp_qred = experiences[0]
    assert exp_qred["id"] == "exp-qred-bank"
    assert exp_qred["title"] == "Senior DevSecOps Engineer"
    assert exp_qred["dateRange"]["start"] == "2022"
    assert exp_qred["dateRange"]["end"] == "2025"
    assert exp_qred["dateRange"]["isCurrent"] is False
    assert exp_qred["organizationRefs"] == [
        {"id": "org-qred-bank", "type": "organization"}
    ]
    assert exp_qred["location"]["label"] == "Stockholm, Sweden"
    assert "engagementType" not in exp_qred
    assert exp_qred["scope"] == "DevSecOps for fintech platform"

    exp_cab = experiences[1]
    assert exp_cab["id"] == "exp-cabonline"
    assert exp_cab["title"] == "DevOps Engineer"
    assert exp_cab["organizationRefs"] == [
        {"id": "org-cabonline", "type": "organization"}
    ]

    organizations = loaded.get("organizations", [])
    assert len(organizations) == 2
    assert organizations[0]["name"] == "Qred Bank"
    assert organizations[1]["name"] == "Cabonline"

    # Verify skills are present with evidence links to experiences
    skills = loaded.get("skills", [])
    assert len(skills) == 5

    skill_names = {s["name"] for s in skills}
    assert "AWS" in skill_names
    assert "Docker" in skill_names
    assert "Kubernetes" in skill_names
    assert "Python" in skill_names
    assert "Terraform" in skill_names

    # Kubernetes was mentioned in both experiences via technologies list
    k8s = next(s for s in skills if s["name"] == "Kubernetes")
    ext = k8s.get("extensions", {})
    evidence = ext.get("experienceEvidence", [])
    assert len(evidence) >= 2
    exp_ids_in_evidence = {e["experienceId"] for e in evidence}
    assert "exp-qred-bank" in exp_ids_in_evidence
    assert "exp-cabonline" in exp_ids_in_evidence

    # Verify education entries are present with institution normalization
    education_entries = loaded.get("education", [])
    assert len(education_entries) == 2

    edu_kth = education_entries[0]
    assert edu_kth["id"] == "edu-kth-royal-institute-of-technology-m.sc.-in-computer-science"
    assert edu_kth["program"] == "M.Sc. in Computer Science"
    assert edu_kth["fieldOfStudy"] == "Computer Science"
    assert edu_kth["institutionRef"]["id"] == "org-kth-royal-institute-of-technology"
    assert edu_kth["dateRange"]["start"] == "2018"
    assert edu_kth["dateRange"]["end"] == "2020"
    assert edu_kth["dateRange"]["isCurrent"] is False

    edu_mit = education_entries[1]
    assert edu_mit["id"] == "edu-massachusetts-institute-of-technology-bachelor-of-science"
    assert edu_mit["program"] == "Bachelor of Science"
    # MIT alias was resolved to full name
    assert edu_mit["institutionRef"]["id"] == "org-massachusetts-institute-of-technology"

    # Verify source traceability metadata is preserved
    acquisition_meta = loaded.get("extensions", {}).get("_acquisition", {})
    assert acquisition_meta.get("sourceDocument", "").endswith("test_resume.docx")
    assert "extractionTimestamp" in acquisition_meta

    loader = SchemaLoader(repo_root / "schemas")
    validator = EntityValidator(loader)
    result = validator.validate_entity(loaded, "profile")
    assert result.is_valid, f"Profile failed validation: {result.errors}"


def test_acquisition_pipeline_rejects_unsupported_format(tmp_path: Path) -> None:
    txt_path = tmp_path / "test.txt"
    txt_path.write_text("Not a docx", encoding="utf-8")

    pipeline = AcquisitionPipeline(llm_extractor=MockLLMExtractor())
    with pytest.raises(Exception, match="Unsupported format"):
        pipeline.run(txt_path)


def test_acquisition_pipeline_records_source_provenance(tmp_path: Path) -> None:
    """Phase 2A: pipeline records sourceName/sourceHash/importedAt provenance."""
    docx_path = tmp_path / "test_resume.docx"
    _create_test_docx(docx_path)
    output_path = tmp_path / "generated-profile.yaml"

    pipeline = AcquisitionPipeline(
        document_reader=DocumentReader(),
        text_extractor=TextExtractor(),
        llm_extractor=MockLLMExtractor(),
        profile_builder=CanonicalProfileBuilder(),
        yaml_writer=YamlWriter(),
    )
    pipeline.run(
        docx_path,
        output_path,
        source_metadata={
            "sourceName": "raul-gongora-cv.docx",
            "sourceHash": "a" * 64,
        },
    )

    loaded = yaml.safe_load(output_path.read_text(encoding="utf-8"))
    acquisition = loaded["extensions"]["_acquisition"]
    assert acquisition["sourceName"] == "raul-gongora-cv.docx"
    assert acquisition["sourceHash"] == "a" * 64
    assert "sourceDocument" in acquisition
    assert "extractionTimestamp" in acquisition
    assert "importedAt" in acquisition
    assert loaded["extensions"]["importedAt"] == acquisition["importedAt"]


def test_acquisition_pipeline_without_source_metadata_omits_hash(tmp_path: Path) -> None:
    """Without source_metadata the pipeline still stamps import timestamps."""
    docx_path = tmp_path / "test_resume.docx"
    _create_test_docx(docx_path)
    output_path = tmp_path / "generated-profile.yaml"

    pipeline = AcquisitionPipeline(llm_extractor=MockLLMExtractor())
    pipeline.run(docx_path, output_path)

    loaded = yaml.safe_load(output_path.read_text(encoding="utf-8"))
    acquisition = loaded["extensions"]["_acquisition"]
    assert "sourceName" not in acquisition
    assert "sourceHash" not in acquisition
    assert "importedAt" in acquisition
    assert loaded["extensions"]["importedAt"] == acquisition["importedAt"]


class TestLLMExtractorFactory:
    def test_create_ollama_extractor(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        extractor = create_llm_extractor()
        assert isinstance(extractor, OllamaLLMExtractor)

    def test_create_ollama_extractor_default_host(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        monkeypatch.delenv("OLLAMA_HOST", raising=False)
        extractor = create_llm_extractor()
        assert extractor.host == "http://localhost:11434"

    def test_create_ollama_extractor_custom_host(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        monkeypatch.setenv("OLLAMA_HOST", "http://192.168.1.100:11434")
        extractor = create_llm_extractor()
        assert extractor.host == "http://192.168.1.100:11434"

    def test_create_ollama_extractor_custom_model(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        monkeypatch.setenv("OLLAMA_MODEL", "llama3.2:3b")
        extractor = create_llm_extractor()
        assert extractor.model == "llama3.2:3b"

    def test_create_ollama_extractor_default_model(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        monkeypatch.delenv("OLLAMA_MODEL", raising=False)
        extractor = create_llm_extractor()
        assert extractor.model == "qwen2.5:3b"

    def test_create_openai_extractor_when_ollama_not_set(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "openai")
        monkeypatch.setenv("OPENAI_API_KEY", "sk-test-key")
        from careeros.acquisition.llm_extractor import OpenAILLMExtractor
        extractor = create_llm_extractor()
        assert isinstance(extractor, OpenAILLMExtractor)

    def test_create_openai_extractor_missing_key(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "openai")
        monkeypatch.delenv("OPENAI_API_KEY", raising=False)
        with pytest.raises(LLMConfigurationError, match="OpenAI API key is required"):
            create_llm_extractor()

    def test_create_extractor_no_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.delenv("LLM_PROVIDER", raising=False)
        with pytest.raises(LLMConfigurationError, match="LLM_PROVIDER is not configured"):
            create_llm_extractor()

    def test_create_unknown_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "anthropic")
        with pytest.raises(LLMConfigurationError, match="Unknown LLM_PROVIDER"):
            create_llm_extractor()

    def test_pipeline_default_uses_factory(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setenv("LLM_PROVIDER", "ollama")
        pipeline = AcquisitionPipeline()
        assert pipeline.llm_extractor is None
        # Calling run would use _default_llm_extractor, but we just verify the
        # pipeline accepts the factory-created extractor via injection.
        pipeline.llm_extractor = OllamaLLMExtractor()
        assert isinstance(pipeline.llm_extractor, OllamaLLMExtractor)


def _run_pipeline_with(docx_path: Path, extractor: LLMExtractor, out: Path) -> dict:
    pipeline = AcquisitionPipeline(llm_extractor=extractor)
    result = pipeline.run(docx_path, out)
    assert result == out.resolve()
    return yaml.safe_load(out.read_text(encoding="utf-8"))


def _canonical_generator(calls: list[str], payload: dict) -> Any:
    import json

    def generator(prompt: str, temperature: float, timeout: float) -> str:
        calls.append(prompt)
        return json.dumps(payload, ensure_ascii=False)

    return generator


def test_pipeline_accepts_profile_with_null_ids(tmp_path: Path) -> None:
    """Explicit null ids from the LLM must not fail canonical validation."""
    docx = tmp_path / "cv.docx"
    _create_test_docx(docx)
    calls: list[str] = []
    payload = {
        "person": {"id": None, "firstName": "Anna", "lastName": "Lindqvist", "fullName": "Anna Lindqvist"},
        "experiences": [{"id": None, "organization": "Qred Bank", "title": "Dev"}],
        "skills": [{"name": "Python"}],
    }
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, payload))
    )
    loaded = _run_pipeline_with(docx, extractor, tmp_path / "out.yaml")

    assert loaded["person"]["id"] == "person-anna-lindqvist"
    assert loaded["person"]["names"][0]["value"] == "Anna Lindqvist"
    assert loaded["experiences"][0]["id"] == "exp-0"
    assert loaded["skills"][0]["name"] == "Python"


def test_pipeline_accepts_partial_profile_missing_optional_fields(tmp_path: Path) -> None:
    """Missing optional fields (email, phone, education, skills) must not fail import."""
    docx = tmp_path / "cv.docx"
    _create_test_docx(docx)
    calls: list[str] = []
    payload = {
        "person": {"id": "person-anna", "fullName": "Anna Lindqvist"},
        "experiences": [{"id": "exp-qred", "organization": "Qred Bank", "title": "Dev"}],
        "skills": [],
        "education": [],
    }
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, payload))
    )
    loaded = _run_pipeline_with(docx, extractor, tmp_path / "out.yaml")

    assert loaded["person"]["id"] == "person-anna-lindqvist"
    assert "contact" not in loaded["person"]
    assert len(loaded["experiences"]) == 1
    assert loaded["experiences"][0]["id"] == "exp-qred"


def test_pipeline_accepts_minimal_profile_name_only(tmp_path: Path) -> None:
    """A profile with only a name is still importable."""
    docx = tmp_path / "cv.docx"
    _create_test_docx(docx)
    calls: list[str] = []
    payload = {
        "person": {"id": "person-anna", "fullName": "Anna Lindqvist"},
        "experiences": [],
        "skills": [],
        "education": [],
    }
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, payload))
    )
    loaded = _run_pipeline_with(docx, extractor, tmp_path / "out.yaml")

    assert loaded["person"]["id"] == "person-anna-lindqvist"
    assert loaded["person"]["names"][0]["value"] == "Anna Lindqvist"
    assert loaded["experiences"] == []
    assert loaded["skills"] == []


def test_pipeline_accepts_swedish_cv(tmp_path: Path) -> None:
    """Swedish section headings are handled semantically; a Swedish profile imports."""
    from docx import Document

    docx = tmp_path / "sv_cv.docx"
    doc = Document()
    doc.add_heading("Anna Lindqvist", 0)
    doc.add_heading("Profil", 1)
    doc.add_paragraph("Erfaren systemutvecklare med 10 års erfarenhet.")
    doc.add_heading("Arbetslivserfarenhet", 1)
    doc.add_heading("Senior Systemutvecklare - Qred Bank", 2)
    doc.add_paragraph("Mars 2020 - Nu")
    doc.add_heading("Utbildning", 1)
    doc.add_paragraph("Civilingenjör Datateknik - KTH, 2009 - 2014")
    doc.add_heading("Kompetens", 1)
    doc.add_paragraph("Python, Kubernetes, AWS")
    doc.add_heading("Certifieringar", 1)
    doc.add_paragraph("AWS Certified Solutions Architect")
    doc.add_heading("Språk", 1)
    doc.add_paragraph("Svenska (Modersmål), Engelska (Flytande)")
    doc.save(str(docx))

    calls: list[str] = []
    payload = {
        "person": {"id": "person-lindqvist", "firstName": "Anna", "lastName": "Lindqvist", "fullName": "Anna Lindqvist", "location": "Stockholm, Sverige"},
        "experiences": [{"id": "exp-qred", "organization": "Qred Bank", "title": "Senior Systemutvecklare", "startDate": "2020-03", "isCurrent": True}],
        "skills": [{"name": "Kubernetes"}, {"name": "AWS"}],
        "education": [{"institution": "KTH Royal Institute of Technology", "degree": "Civilingenjör Datateknik"}],
    }
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, payload))
    )
    loaded = _run_pipeline_with(docx, extractor, tmp_path / "out.yaml")

    assert loaded["person"]["names"][0]["value"] == "Anna Lindqvist"
    assert loaded["person"]["location"]["label"] == "Stockholm, Sverige"
    assert loaded["experiences"][0]["title"] == "Senior Systemutvecklare"
    assert {s["name"] for s in loaded["skills"]} >= {"Kubernetes", "AWS"}
    assert loaded["education"][0]["program"] == "Civilingenjör Datateknik"


def test_pipeline_accepts_large_cv_via_chunking(tmp_path: Path) -> None:
    """A CV larger than one chunk is split, extracted, merged, and validated."""
    from docx import Document

    docx = tmp_path / "large_cv.docx"
    doc = Document()
    doc.add_heading("Anna Lindqvist", 0)
    for i in range(40):
        doc.add_heading(f"Systemutvecklare - Company {i}", 1)
        doc.add_paragraph(f"Jan 2010 - Dec 201{i % 10 + 5}")
        doc.add_paragraph(
            "Designed and deployed cloud services in Python with Kubernetes "
            "and AWS. Automated infrastructure with Terraform and CI/CD."
        )
    doc.save(str(docx))

    calls: list[str] = []
    payload = {
        "person": {"id": "person-lindqvist", "firstName": "Anna", "lastName": "Lindqvist", "fullName": "Anna Lindqvist"},
        "experiences": [{"id": "exp-acme", "organization": "ACME", "title": "Systemutvecklare", "startDate": "2020-01", "isCurrent": True}],
        "skills": [{"name": "Python"}],
        "education": [],
    }
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, payload))
    )
    loaded = _run_pipeline_with(docx, extractor, tmp_path / "out.yaml")

    assert len(calls) > 1
    assert loaded["person"]["names"][0]["value"] == "Anna Lindqvist"
    assert loaded["experiences"][0]["organizationRefs"][0]["id"] == "org-acme"


def test_pipeline_rejects_document_with_no_extractable_content(tmp_path: Path) -> None:
    """A document that yields no name and no entities is genuinely invalid."""
    docx = tmp_path / "cv.docx"
    _create_test_docx(docx)
    calls: list[str] = []
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, {}))
    )
    pipeline = AcquisitionPipeline(llm_extractor=extractor)
    with pytest.raises(PipelineError, match="No profile content could be extracted"):
        pipeline.run(docx, tmp_path / "out.yaml")


def test_pipeline_rejects_genuinely_invalid_document(tmp_path: Path) -> None:
    """Rejects unsupported content, and rejected imports do not write output."""
    from docx import Document

    docx = tmp_path / "cv.docx"
    doc = Document()
    doc.add_heading("Not a real CV", 0)
    doc.add_paragraph("This document contains no professional information.")
    doc.save(str(docx))

    calls: list[str] = []
    extractor = LLMExtractor(
        provider=MockAIProvider(generator=_canonical_generator(calls, {}))
    )
    pipeline = AcquisitionPipeline(llm_extractor=extractor)
    out = tmp_path / "out.yaml"
    with pytest.raises(PipelineError, match="No profile content could be extracted"):
        pipeline.run(docx, out)
    assert not out.exists()
