from pathlib import Path

import pytest

from careeros.acquisition.document_reader import DocumentReader, DocumentReadError


def test_read_docx_returns_text(tmp_path: Path) -> None:
    docx_path = _create_test_docx(tmp_path, "Hello World\nThis is a resume.")
    reader = DocumentReader()
    text = reader.read(docx_path)

    assert "Hello World" in text
    assert "resume" in text


def test_read_docx_includes_table_content(tmp_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    docx_path = tmp_path / "test_table.docx"
    doc = Document()
    doc.add_paragraph("Skills:")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Advanced"
    doc.save(str(docx_path))

    reader = DocumentReader()
    text = reader.read(docx_path)

    assert "Skill | Level" in text
    assert "Python | Advanced" in text


def test_read_raises_for_missing_file() -> None:
    reader = DocumentReader()
    with pytest.raises(DocumentReadError, match="not found"):
        reader.read("/nonexistent/path.docx")


def test_read_raises_for_unsupported_format(tmp_path: Path) -> None:
    reader = DocumentReader()
    path = tmp_path / "test.txt"
    path.write_text("not a docx", encoding="utf-8")
    with pytest.raises(DocumentReadError, match="Unsupported format"):
        reader.read(path)


def test_read_raises_for_invalid_docx(tmp_path: Path) -> None:
    reader = DocumentReader()
    path = tmp_path / "invalid.docx"
    path.write_bytes(b"not a valid docx file")
    with pytest.raises(DocumentReadError):
        reader.read(path)


def _create_test_docx(tmp_path: Path, text: str) -> Path:
    from docx import Document

    path = tmp_path / "test_resume.docx"
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(path))
    return path
