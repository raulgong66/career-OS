"""Shared Markdown -> DOCX rendering helpers for CareerOS artifact generators.

All artifact DOCX generators render by streaming Markdown into ``python-docx``
using the same minimal subset: ``# `` / ``## `` headings, ``- `` / ``• ``
bullets, ``_italic_`` lines, and plain paragraphs with ``**bold**`` inline
runs. This module is the single source of truth for that mapping so new
generators (CV, letters, preparation guides, future artifacts) reuse the
behavior instead of duplicating it.
"""

from __future__ import annotations

from typing import TYPE_CHECKING


if TYPE_CHECKING:
    from docx.document import Document as _Document
    from docx.text.paragraph import Paragraph as _Paragraph


def add_markdown_line(document: "_Document", line: str) -> None:
    """Add a Markdown-derived line to a ``python-docx`` document.

    Supports the Markdown subset emitted by CareerOS generators:

    - ``# heading``   -> Heading 1 (level 0)
    - ``## heading``  -> Heading 2 (level 1)
    - ``- item`` / ``• item`` -> List Bullet paragraph
    - ``_italic_``    -> Italic paragraph
    - everything else -> Plain paragraph with ``**bold**`` inline runs
    """
    if not line:
        return
    if line.startswith("# "):
        document.add_heading(line[2:], level=0)
        return
    if line.startswith("## "):
        document.add_heading(line[3:], level=1)
        return
    if line.startswith("- ") or line.startswith("• "):
        paragraph = document.add_paragraph(style="List Bullet")
        add_inline_text(paragraph, line[2:])
        return
    if line.startswith("_") and line.endswith("_"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line.strip("_"))
        run.italic = True
        return
    paragraph = document.add_paragraph()
    add_inline_text(paragraph, line)


def add_inline_text(paragraph: "_Paragraph", text: str) -> None:
    """Add ``text`` to ``paragraph`` with minimal Markdown ``**bold**`` support."""
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = index % 2 == 1
