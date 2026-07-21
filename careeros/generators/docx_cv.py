"""DOCX CV generator."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from ..exceptions import ValidationError
from ..export_contract import ExportContract
from .markdown_cv import MarkdownCVGenerator


class DocxCVGenerator:
    """Generate a minimal DOCX CV from an export contract."""

    supported_artifact_types = {"CV", "RESUME"}

    def __init__(self) -> None:
        """Create a DOCX generator backed by the Markdown CV structure."""
        self.markdown_generator = MarkdownCVGenerator()

    def generate(self, contract: ExportContract) -> bytes:
        """Generate DOCX bytes using only the provided export contract."""
        artifact_type = contract.artifact_type.upper()
        if artifact_type not in self.supported_artifact_types:
            raise ValidationError(f"Unsupported artifact type for DOCX CV: {contract.artifact_type}")

        markdown = self.markdown_generator.generate(contract)
        document = Document()
        for line in markdown.splitlines():
            self._add_markdown_line(document, line)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _add_markdown_line(self, document: Document, line: str) -> None:
        """Add a Markdown-derived line to a DOCX document."""
        if not line:
            return
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
            return
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            return
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            self._add_inline_text(paragraph, line[2:])
            return
        if line.startswith("_") and line.endswith("_"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.strip("_"))
            run.italic = True
            return
        paragraph = document.add_paragraph()
        self._add_inline_text(paragraph, line)

    @staticmethod
    def _add_inline_text(paragraph, text: str) -> None:
        """Add a line of text with minimal Markdown bold support."""
        parts = text.split("**")
        for index, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = index % 2 == 1
