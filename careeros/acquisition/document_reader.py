from __future__ import annotations

from pathlib import Path

from careeros.exceptions import CareerOSException


class DocumentReadError(CareerOSException):
    pass


class DocumentReader:
    SUPPORTED_EXTENSIONS = {".docx"}

    def read(self, path: str | Path) -> str:
        resolved = Path(path).expanduser().resolve()
        if not resolved.exists():
            raise DocumentReadError(f"File not found: {resolved}")
        suffix = resolved.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise DocumentReadError(
                f"Unsupported format '{suffix}'. Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )
        return self._read_docx(resolved)

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise DocumentReadError(
                "python-docx is required to read DOCX files. Install it with: pip install python-docx"
            )
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise DocumentReadError(f"Failed to open DOCX file: {exc}") from exc

        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                paragraphs.append(row_text)

        return "\n".join(paragraphs)
