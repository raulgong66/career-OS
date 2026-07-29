"""Pure rendering layer to update DOCX CV documents with recommendations."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any, Union

import docx
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from .optimizer import Recommendation


class CVDocumentRenderer:
    """Pure renderer for updating DOCX CV files with optimized content."""

    def __init__(self) -> None:
        """Initialize the renderer."""
        pass

    def apply_recommendations(
        self,
        docx_path: Union[str, Path],
        output_path: Union[str, Path],
        recommendations: list[Recommendation],
    ) -> None:
        """Apply recommended additions to a DOCX document.

        Args:
            docx_path: Path to the input DOCX file.
            output_path: Path where the updated DOCX file should be written.
            recommendations: List of Recommendation objects (with operation='ADD').
        """
        input_path = Path(docx_path).expanduser().resolve()
        save_path = Path(output_path).expanduser().resolve()

        if not input_path.exists():
            raise FileNotFoundError(f"Input DOCX file not found: {input_path}")

        # Load the document
        doc = docx.Document(str(input_path))

        # Filter recommendations by operation type 'ADD'
        additions = [r for r in recommendations if r.operation == "ADD"]
        if not additions:
            doc.save(str(save_path))
            return

        # Group additions by recommendation type
        grouped_additions: dict[str, list[Recommendation]] = {}
        for rec in additions:
            grouped_additions.setdefault(rec.type, []).append(rec)

        # Categorize sections and match them to categories
        categories_map = {
            "skill": {
                "keywords": {"skills", "technical skills", "competencies", "expertise"},
                "section_title": "Skills",
            },
            "experience": {
                "keywords": {"experience", "work history", "employment history", "professional experience"},
                "section_title": "Professional Experience",
            },
            "achievement": {
                "keywords": {"achievements", "accomplishments", "milestones"},
                "section_title": "Achievements",
            },
            "project": {
                "keywords": {"projects", "key projects", "selected projects", "portfolio"},
                "section_title": "Projects",
            },
            "education": {
                "keywords": {"education", "academic history", "qualification"},
                "section_title": "Education",
            },
            "certification": {
                "keywords": {"certifications", "credentials", "licenses", "training"},
                "section_title": "Certifications",
            },
        }

        # Locate section headers in the document
        # Map of category to paragraph index
        headings_found: dict[str, int] = {}
        for idx, p in enumerate(doc.paragraphs):
            text = p.text.lower().strip()
            if not text:
                continue

            # Heuristics for finding headings
            is_likely_heading = (
                p.style.name.startswith("Heading")
                or text.isupper()
                or any(run.bold for run in p.runs)
            )

            if is_likely_heading:
                for cat, config in categories_map.items():
                    # Exact or substring match of the keyword
                    if any(kw in text for kw in config["keywords"]):
                        headings_found[cat] = idx
                        break

        # Process each category
        for cat, recs in grouped_additions.items():
            if not recs:
                continue

            # Determine where to insert the new items
            if cat in headings_found:
                # Find the last paragraph of the section before the next heading starts
                start_idx = headings_found[cat]
                section_paragraphs = []
                for i in range(start_idx + 1, len(doc.paragraphs)):
                    para = doc.paragraphs[i]
                    txt = para.text.lower().strip()
                    # A new section heading is identified by styling or all-caps
                    is_next_heading = (
                        para.style.name.startswith("Heading")
                        or (txt.isupper() and len(txt) > 2)
                    )
                    if is_next_heading:
                        break
                    section_paragraphs.append(para)

                # Anchor paragraph to insert after
                anchor = doc.paragraphs[start_idx]
                style_to_use = None

                if section_paragraphs:
                    # Find a non-empty paragraph to copy style/bullet formatting
                    for p in reversed(section_paragraphs):
                        if p.text.strip():
                            anchor = p
                            # Check if the paragraph has a list-bullet style
                            if p.style and ("bullet" in p.style.name.lower() or "list" in p.style.name.lower()):
                                style_to_use = p.style
                            break
                else:
                    # Check if the heading itself has a next-paragraph style
                    pass

                # If no list style was found, default to 'List Bullet' if we are inserting list items
                if not style_to_use:
                    try:
                        style_to_use = doc.styles['List Bullet']
                    except KeyError:
                        style_to_use = None

                # Insert recommendations sequentially after the anchor
                for rec in recs:
                    formatted_text = self._format_recommendation(rec)
                    new_p = self._insert_paragraph_after(anchor, formatted_text, style_to_use)
                    anchor = new_p

            else:
                # Section heading not found: Append heading and items to the end of the document
                config = categories_map.get(cat, {"section_title": cat.capitalize()})
                
                # Append heading
                heading_p = doc.add_paragraph()
                try:
                    heading_p.style = doc.styles['Heading 2']
                except KeyError:
                    pass
                run = heading_p.add_run(config["section_title"])
                run.bold = True
                
                # Append recommendations
                try:
                    bullet_style = doc.styles['List Bullet']
                except KeyError:
                    bullet_style = None

                for rec in recs:
                    formatted_text = self._format_recommendation(rec)
                    p = doc.add_paragraph(style=bullet_style)
                    r = p.add_run(formatted_text)
                    r.font.color.rgb = RGBColor(0, 0, 255)  # Highlight in blue

        # Save output document
        save_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(save_path))

    def _insert_paragraph_after(self, paragraph: Paragraph, text: str, style: Any = None) -> Paragraph:
        """Insert a new paragraph directly after the given paragraph."""
        p_element = OxmlElement("w:p")
        paragraph._p.addnext(p_element)
        new_paragraph = Paragraph(p_element, paragraph._parent)

        if style:
            new_paragraph.style = style
        elif paragraph.style:
            new_paragraph.style = paragraph.style

        # Add run and format color in blue
        run = new_paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)

        return new_paragraph

    def _format_recommendation(self, rec: Recommendation) -> str:
        """Format the recommendation into a friendly text string with backing evidence details."""
        details = rec.details
        evidence_str = ""

        if rec.evidence:
            ev_titles = [ev.get("title") or ev.get("id") for ev in rec.evidence if ev.get("title") or ev.get("id")]
            if ev_titles:
                evidence_str = f" [Verified by: {', '.join(ev_titles)}]"

        if rec.type == "skill":
            return f"{rec.display_name}{evidence_str}"
        
        elif rec.type == "experience":
            date_range = details.get("dateRange", {})
            start = date_range.get("start", "")
            end = date_range.get("end", "")
            date_str = f" ({start} - {end})" if start else ""
            scope = details.get("scope", "")
            scope_str = f" - {scope}" if scope else ""
            return f"{rec.display_name}{date_str}{scope_str}{evidence_str}"
        
        elif rec.type == "achievement":
            return f"{rec.display_name}{evidence_str}"
        
        elif rec.type == "project":
            desc = details.get("description", "")
            desc_str = f" - {desc}" if desc else ""
            return f"{rec.display_name}{desc_str}{evidence_str}"
        
        elif rec.type == "education":
            date_range = details.get("dateRange", {})
            start = date_range.get("start", "")
            end = date_range.get("end", "")
            date_str = f" ({start} - {end})" if start else ""
            return f"{rec.display_name}{date_str}{evidence_str}"
        
        elif rec.type == "certification":
            cred_id = details.get("credentialId", "")
            cred_str = f" (Credential ID: {cred_id})" if cred_id else ""
            return f"{rec.display_name}{cred_str}{evidence_str}"

        return f"{rec.display_name}{evidence_str}"
